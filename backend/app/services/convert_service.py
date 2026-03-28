import os
from typing import List
import pdfplumber
import io
import zipfile
from pypdf import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from docx import Document
from pdf2image import convert_from_path
from PIL import Image
from openpyxl import Workbook

from app.utils.file_handler import create_temp_file
from app.models.schemas import PdfToWordOptions, WordToPdfOptions


def pdf_to_word(pdf_path: str, options: PdfToWordOptions) -> str:
    # Converts PDF text into a DOCX document
    out_path, out_id = create_temp_file(".docx")
    doc = Document()
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if options.include_page_breaks and i > 1:
                doc.add_page_break()
            doc.add_paragraph(text)
    doc.save(out_path)
    return out_id


def word_to_pdf(docx_path: str, options: WordToPdfOptions) -> str:
    # Renders DOCX text to a simple PDF using reportlab
    out_path, out_id = create_temp_file(".pdf")
    doc = Document(docx_path)
    c = canvas.Canvas(out_path, pagesize=A4)
    width, height = A4
    x_margin = 1 * inch
    y_margin = 1 * inch
    y = height - y_margin
    c.setFont(options.font_name, options.font_size)
    line_height = options.font_size * options.line_height
    for p in doc.paragraphs:
        text = p.text or ""
        for line in text.split("\n"):
            if y < y_margin + line_height:
                c.showPage()
                c.setFont(options.font_name, options.font_size)
                y = height - y_margin
            c.drawString(x_margin, y, line)
            y -= line_height
        y -= line_height * 0.5
    c.save()
    return out_id


def jpg_to_pdf(image_paths: List[str]) -> str:
    # Merges one or more images into a single PDF
    out_path, out_id = create_temp_file(".pdf")
    images = [Image.open(p).convert("RGB") for p in image_paths]
    if not images:
        raise ValueError("No images provided")
    first, rest = images[0], images[1:]
    first.save(out_path, save_all=True, append_images=rest)
    return out_id


def pdf_to_jpg(pdf_path: str) -> List[str]:
    # Converts PDF pages to JPG images
    pages = convert_from_path(pdf_path)
    outputs: List[str] = []
    for i, img in enumerate(pages, start=1):
        out_path, out_id = create_temp_file(".jpg")
        img.save(out_path, "JPEG")
        outputs.append(out_id)
    return outputs


def pdf_to_excel(pdf_path: str) -> str:
    # Extracts tables from PDF into an XLSX workbook
    out_path, out_id = create_temp_file(".xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Page1"
    with pdfplumber.open(pdf_path) as pdf:
        sheet_index = 1
        for page in pdf.pages:
            tables = page.extract_tables() or []
            if sheet_index == 1:
                target_ws = ws
            else:
                target_ws = wb.create_sheet(title=f"Page{sheet_index}")
            row_idx = 1
            if tables:
                for tbl in tables:
                    for row in tbl:
                        target_ws.append([cell if cell is not None else "" for cell in row])
                        row_idx += 1
                    row_idx += 1
            else:
                text = page.extract_text() or ""
                for line in text.splitlines():
                    target_ws.append([line])
            sheet_index += 1
    wb.save(out_path)
    return out_id

def pdf_to_ppt(pdf_path: str) -> str:
    out_path, out_id = create_temp_file(".pptx")
    pages = convert_from_path(pdf_path)
    slides = len(pages)
    if slides == 0:
        raise ValueError("No pages provided")
    images: List[bytes] = []
    for img in pages:
        b = io.BytesIO()
        img.save(b, format="JPEG")
        images.append(b.getvalue())
    ct = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">',
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>',
        '<Default Extension="xml" ContentType="application/xml"/>',
        '<Default Extension="jpeg" ContentType="image/jpeg"/>',
        '<Default Extension="jpg" ContentType="image/jpeg"/>',
        '<Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>',
    ]
    for i in range(1, slides + 1):
        ct.append(f'<Override PartName="/ppt/slides/slide{i}.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>')
    ct.append("</Types>")
    ct_xml = "\n".join(ct)
    rels_xml = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/>
</Relationships>
"""
    pres_rels = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">',
    ]
    for i in range(1, slides + 1):
        pres_rels.append(f'<Relationship Id="rId{i}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide{i}.xml"/>')
    pres_rels.append("</Relationships>")
    pres_rels_xml = "\n".join(pres_rels)
    sld_ids = ['<p:sldIdLst>']
    for i in range(1, slides + 1):
        sld_ids.append(f'<p:sldId id="{255 + i}" r:id="rId{i}"/>')
    sld_ids.append("</p:sldIdLst>")
    pres_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
{sld_ids[0]}
{''.join(sld_ids[1:-1])}
{sld_ids[-1]}
<p:sldSz cx="9144000" cy="6858000" type="screen4x3"/>
<p:notesSz cx="6858000" cy="9144000"/>
</p:presentation>
'''
    def slide_xml(i: int) -> str:
        return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld>
    <p:spTree>
      <p:nvGrpSpPr>
        <p:cNvPr id="1" name=""/>
        <p:cNvGrpSpPr/>
        <p:nvPr/>
      </p:nvGrpSpPr>
      <p:grpSpPr><a:xfrm/></p:grpSpPr>
      <p:pic>
        <p:nvPicPr>
          <p:cNvPr id="2" name="Picture {i}"/>
          <p:cNvPicPr/>
          <p:nvPr/>
        </p:nvPicPr>
        <p:blipFill>
          <a:blip r:embed="rId1"/>
          <a:stretch><a:fillRect/></a:stretch>
        </p:blipFill>
        <p:spPr>
          <a:xfrm><a:off x="0" y="0"/><a:ext cx="9144000" cy="6858000"/></a:xfrm>
        </p:spPr>
      </p:pic>
    </p:spTree>
  </p:cSld>
</p:sld>
'''
    def slide_rels_xml(i: int) -> str:
        return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="../media/image{i}.jpeg"/>
</Relationships>
'''
    with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", ct_xml)
        z.writestr("_rels/.rels", rels_xml)
        z.writestr("ppt/presentation.xml", pres_xml)
        z.writestr("ppt/_rels/presentation.xml.rels", pres_rels_xml)
        for i, b in enumerate(images, start=1):
            z.writestr(f"ppt/slides/slide{i}.xml", slide_xml(i))
            z.writestr(f"ppt/slides/_rels/slide{i}.xml.rels", slide_rels_xml(i))
            z.writestr(f"ppt/media/image{i}.jpeg", b)
    return out_id
